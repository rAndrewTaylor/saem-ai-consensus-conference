"""Word document rendering for the Round 1 report.

Composes the .docx from pre-aggregated data, figure PNG bytes, and
AI-synthesized prose. Returns the document as bytes — caller decides
whether to write to disk, stream to a download endpoint, etc.

The document structure follows Round_1_Analysis_Plan.md §2:
  Cover
  A — Process overview
  B — Headline numbers
  C — Per-WG deep dive (one section per WG)
  D — Cross-WG analysis  (later phase; left as a stub here)
  E — Round 2 plan       (later phase; stub)
  Appendices             (later phase; stub)
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Optional

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

# Locked decision: every gated page / DOCX gets a watermark line
WATERMARK = "Pre-publication — SAEM 2026 working draft, do not distribute"


def render(
    *,
    bundle,
    q_stats: pd.DataFrame,
    wg_summary: pd.DataFrame,
    overall: dict,
    per_wg_figures: dict[int, dict[str, bytes]],
    cross_figures: Optional[dict[str, bytes]] = None,
    pillar_tags: Optional[pd.DataFrame] = None,
    cross_cutting_tags: Optional[pd.DataFrame] = None,
    theme_clusters: Optional[list] = None,
    similarity_data: Optional[tuple] = None,  # (qids, np.ndarray) for D.2 overlap table
    output_path: Optional[str] = None,
) -> bytes:
    """Build the report .docx and return its bytes.

    Args:
      bundle:   ReportBundle (for snapshot timestamp + suggestions)
      q_stats:  per_question_stats output
      wg_summary: per_wg_summary output
      overall:  overall_summary output
      per_wg_figures: { wg_id: {"F1": png_bytes, "F2": png_bytes, ...}, ... }
      output_path: optional path to also save to disk
    """
    doc = Document()
    _configure_styles(doc)

    _cover(doc, overall)
    _section_a(doc, overall, wg_summary)
    _section_b(doc, q_stats, wg_summary, overall)
    for _, w in wg_summary.iterrows():
        _section_c_wg(
            doc,
            wg=w,
            q_stats=q_stats[q_stats["wg_id"] == w["wg_id"]],
            figures=per_wg_figures.get(int(w["wg_id"]), {}),
            suggestions=bundle.suggestions[bundle.suggestions["wg_id"] == w["wg_id"]]
                if not bundle.suggestions.empty else pd.DataFrame(),
        )
    _section_d(
        doc,
        cross_figures=cross_figures or {},
        pillar_tags=pillar_tags,
        cross_cutting_tags=cross_cutting_tags,
        theme_clusters=theme_clusters or [],
        similarity_data=similarity_data,
        questions=bundle.questions,
    )
    _section_e_stub(doc)
    _appendix_stub(doc)

    buf = io.BytesIO()
    doc.save(buf)
    if output_path:
        with open(output_path, "wb") as f:
            f.write(buf.getvalue())
    return buf.getvalue()


# --- Styling --------------------------------------------------------------

def _configure_styles(doc: "Document") -> None:
    # Body
    body = doc.styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    # Title
    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor(0x0C, 0x23, 0x40)
    title.font.bold = True
    # Headings
    for h, sz in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        s = doc.styles[h]
        s.font.name = "Calibri"
        s.font.size = Pt(sz)
        s.font.color.rgb = RGBColor(0x0C, 0x23, 0x40)
        s.font.bold = True


def _watermark(doc: "Document") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(WATERMARK)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)


# --- Cover ----------------------------------------------------------------

def _cover(doc: "Document", overall: dict) -> None:
    snapshot = overall.get("snapshot_at", datetime.utcnow().isoformat())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SAEM 2026 AI Consensus Conference")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Round 1 Inter-Round Report")
    tr.font.size = Pt(24)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x0C, 0x23, 0x40)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Modified Delphi · Pairwise comparison · AI synthesis")
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    _watermark(doc)

    # Quick summary table
    rows = [
        ("Snapshot", snapshot),
        ("Working groups", overall.get("n_working_groups", "—")),
        ("Active questions", overall.get("n_questions_total", "—")),
        ("R1 responses", overall.get("n_r1_responses_total", "—")),
        ("Pairwise votes", overall.get("n_pairwise_votes_total", "—")),
        ("Confirmed / gray / removed / open",
            f"{overall.get('n_confirmed_total', 0)} / "
            f"{overall.get('n_gray_total', 0)} / "
            f"{overall.get('n_removed_total', 0)} / "
            f"{overall.get('n_open_total', 0)}"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = str(v)
        for c in (table.cell(i, 0), table.cell(i, 1)):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()


# --- Section A — Process overview ----------------------------------------

def _section_a(doc: "Document", overall: dict, wg_summary: pd.DataFrame) -> None:
    doc.add_heading("A. Process overview", level=1)
    doc.add_heading("A.1 What happened in Round 1", level=2)
    doc.add_paragraph(
        f"Round 1 of the modified Delphi closed with "
        f"{overall.get('n_r1_responders_total', 0)} of "
        f"{overall.get('n_invited_total', 0)} invited members responding "
        f"({_pct(overall.get('n_r1_responders_total', 0), overall.get('n_invited_total', 0))} response rate). "
        f"Across {overall.get('n_working_groups', 5)} working groups, "
        f"{overall.get('n_r1_responses_total', 0)} disposition+importance ratings were submitted on "
        f"{overall.get('n_questions_total', 0)} active questions. "
        f"Pairwise comparison ran in parallel and is still accepting votes — "
        f"{overall.get('n_pairwise_votes_total', 0)} votes from "
        f"{overall.get('n_pairwise_voters_total', 0)} unique voters as of this snapshot."
    )
    doc.add_paragraph(
        f"Free-text feedback: {overall.get('n_comments_total', 0)} comments and "
        f"{overall.get('n_suggestions_total', 0)} new-question suggestions. "
        f"Comments are AI-synthesized into themes per question (Section C); "
        f"suggestions are categorized and reviewed (Appendix 2)."
    )
    low_pw = overall.get("wgs_low_pairwise_confidence", [])
    if low_pw:
        p = doc.add_paragraph()
        p.add_run("Note: ").bold = True
        p.add_run(
            f"WG {', '.join(str(n) for n in low_pw)} pairwise data is below "
            f"the confidence threshold (<8 unique voters or <100 pairs). "
            f"Pairwise figures for those groups are shown with wide confidence "
            f"intervals; treat the rankings as preliminary."
        )

    doc.add_heading("A.2 Methods (recap)", level=2)
    doc.add_paragraph(
        "The full methodology lives in AI_Enhanced_Consensus_Methodology.md. "
        "In short: each question receives a disposition vote (Include / Modify / "
        "Exclude) and an importance rating (1–9 Likert) from each respondent. "
        "Concurrently, participants compare pairs of questions head-to-head; "
        "scores follow a Laplace-smoothed Bradley-Terry model. AI synthesis "
        "(Claude Opus, with full audit logging) clusters comments into themes "
        "and suggests question revisions, all subject to human review by the "
        "co-leads before incorporation into Round 2."
    )

    doc.add_heading("A.3 How to read this report", level=2)
    doc.add_paragraph(
        "Each question is assigned a bucket based on Round 1 results:"
    )
    bullets = doc.add_paragraph()
    bullets.add_run(
        "Confirmed — ≥80% include AND mean importance ≥7. Strong consensus.\n"
        "Gray — 60–79% include, OR ≥80% include with importance <7. Productive disagreement; primary target for Round 2 deliberation.\n"
        "Removed — ≥80% exclude (or <40% include with low importance). Strong consensus to drop.\n"
        "Open — everything else. Heads to Round 2 unchanged for re-rating."
    )
    doc.add_paragraph(
        "Co-leads make the final call per question — buckets are decision aids, "
        "not automatic actions."
    )


# --- Section B — Headline numbers ----------------------------------------

def _section_b(doc: "Document", q_stats: pd.DataFrame, wg_summary: pd.DataFrame, overall: dict) -> None:
    doc.add_heading("B. Headline numbers", level=1)
    doc.add_heading("B.1 Participation by working group", level=2)

    headers = ["WG", "Pillar", "Members", "R1 resp.", "Resp. rate", "Median Qs/resp.",
               "PW voters", "PW votes", "Comments", "Suggestions"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for _, w in wg_summary.iterrows():
        row = table.add_row().cells
        row[0].text = f"WG{int(w['wg_number'])} — {w['short_name']}"
        row[1].text = str(w["pillar"])
        row[2].text = str(int(w["n_invited"]))
        row[3].text = str(int(w["n_r1_responders"]))
        row[4].text = f"{w['response_rate_pct']:.0f}%"
        row[5].text = str(int(w["median_questions_per_responder"]))
        pw_str = f"{int(w['n_pairwise_voters'])}"
        if w["pairwise_low_confidence"]:
            pw_str += "*"
        row[6].text = pw_str
        row[7].text = str(int(w["n_pairwise_votes"]))
        row[8].text = str(int(w["n_comments"]))
        row[9].text = str(int(w["n_suggestions"]))
    p = doc.add_paragraph()
    p.add_run("* Below low-confidence threshold (<8 unique pairwise voters or <100 votes).").italic = True
    p.runs[0].font.size = Pt(8)

    doc.add_heading("B.2 Question outcomes", level=2)
    headers = ["WG", "Active Qs", "Confirmed", "Gray-zone", "Removed", "Open", "Delphi/Pairwise ρ"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs: r.bold = True
    for _, w in wg_summary.iterrows():
        row = t.add_row().cells
        row[0].text = f"WG{int(w['wg_number'])}"
        row[1].text = str(int(w["n_questions"]))
        row[2].text = str(int(w["n_confirmed"]))
        row[3].text = str(int(w["n_gray"]))
        row[4].text = str(int(w["n_removed"]))
        row[5].text = str(int(w["n_open"]))
        rho = w["spearman_rho_delphi_pairwise"]
        row[6].text = "—" if rho is None else f"{rho:.2f}"
    # Totals row
    row = t.add_row().cells
    row[0].text = "All WGs"
    for r in row[0].paragraphs[0].runs: r.bold = True
    row[1].text = str(int(wg_summary["n_questions"].sum()))
    row[2].text = str(int(wg_summary["n_confirmed"].sum()))
    row[3].text = str(int(wg_summary["n_gray"].sum()))
    row[4].text = str(int(wg_summary["n_removed"].sum()))
    row[5].text = str(int(wg_summary["n_open"].sum()))
    rho_all = overall.get("spearman_rho_delphi_pairwise_overall")
    row[6].text = "—" if rho_all is None else f"{rho_all:.2f}"

    doc.add_heading("B.3 Top 10 across working groups (by Delphi importance)", level=2)
    if not q_stats.empty:
        top = (q_stats.dropna(subset=["importance_mean"])
                       .sort_values("importance_mean", ascending=False)
                       .head(10))
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(["WG", "Q-id", "Question", "Importance", "Pairwise"]):
            c = t.rows[0].cells[i]; c.text = h
            for r in c.paragraphs[0].runs: r.bold = True
        for _, q in top.iterrows():
            row = t.add_row().cells
            row[0].text = f"WG{int(q['wg_number'])}"
            row[1].text = f"Q{int(q['question_id'])}"
            short = (q.get("short_text") or q.get("text", "")).strip().replace("\n", " ")
            if len(short) > 90: short = short[:89] + "…"
            row[2].text = short
            row[3].text = f"{q['importance_mean']:.1f}"
            row[4].text = "—" if pd.isna(q.get("pairwise_score")) else f"{q['pairwise_score']:.0f}"

    doc.add_page_break()


# --- Section C — Per-WG deep dive ----------------------------------------

def _section_c_wg(
    doc: "Document",
    *,
    wg: pd.Series,
    q_stats: pd.DataFrame,
    figures: dict[str, bytes],
    suggestions: pd.DataFrame,
) -> None:
    doc.add_heading(f"C.{int(wg['wg_number'])} WG{int(wg['wg_number'])} — {wg['name']}", level=1)
    doc.add_paragraph(
        f"Pillar: {wg['pillar']}. {int(wg['n_invited'])} members invited, "
        f"{int(wg['n_r1_responders'])} responded ({wg['response_rate_pct']:.0f}%), "
        f"{int(wg['n_r1_responses'])} R1 ratings. "
        f"Pairwise: {int(wg['n_pairwise_votes'])} votes from {int(wg['n_pairwise_voters'])} voters"
        f"{' — flagged low-confidence' if wg['pairwise_low_confidence'] else ''}."
    )

    # Bucket distribution
    p = doc.add_paragraph()
    p.add_run("Question outcomes: ").bold = True
    p.add_run(
        f"{int(wg['n_confirmed'])} confirmed · {int(wg['n_gray'])} gray-zone · "
        f"{int(wg['n_removed'])} removed · {int(wg['n_open'])} open."
    )
    rho = wg["spearman_rho_delphi_pairwise"]
    if rho is not None:
        p2 = doc.add_paragraph()
        p2.add_run("Method concordance: ").bold = True
        p2.add_run(f"Delphi importance vs. pairwise score Spearman ρ = {rho:.2f}.")

    # Figures F1-F4
    for fkey, caption in [
        ("F1", "Figure C.{n}.1 — Round 1 disposition by question (sorted by include%)."),
        ("F2", "Figure C.{n}.2 — Round 1 importance ratings (mean ◆, IQR bar)."),
        ("F3", "Figure C.{n}.3 — Pairwise priority ranking with 95% CIs."),
        ("F4", "Figure C.{n}.4 — Delphi vs. pairwise concordance."),
    ]:
        png = figures.get(fkey)
        if not png:
            continue
        doc.add_picture(io.BytesIO(png), width=Inches(6.4))
        cap = doc.add_paragraph(caption.format(n=int(wg["wg_number"])))
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # New questions suggested
    if not suggestions.empty:
        doc.add_heading("New questions suggested", level=3)
        for _, s in suggestions.iterrows():
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(s["text"])
    doc.add_page_break()


# --- Stubs for later phases ----------------------------------------------

def _section_d(
    doc: "Document",
    *,
    cross_figures: dict[str, bytes],
    pillar_tags: Optional[pd.DataFrame],
    cross_cutting_tags: Optional[pd.DataFrame],
    theme_clusters: list,
    similarity_data: Optional[tuple],
    questions: pd.DataFrame,
) -> None:
    doc.add_heading("D. Cross-WG analysis", level=1)
    doc.add_paragraph(
        "Embeddings of every active question (sentence-transformers/"
        "all-MiniLM-L6-v2) feed the similarity network and theme "
        "dendrogram. Pillar tags and cross-cutting topic tags come "
        "from Claude Opus with full audit logging — co-leads should "
        "spot-check 10% of tags before relying on the totals."
    )

    # F6 — similarity network
    if "F6" in cross_figures:
        doc.add_heading("D.1 Question similarity across working groups", level=2)
        doc.add_picture(io.BytesIO(cross_figures["F6"]), width=Inches(6.4))
        cap = doc.add_paragraph(
            "Figure D.1 — Force-directed graph of question similarity. "
            "Nodes are sized by Round 1 importance and colored by WG. "
            "Cross-WG edges are highlighted in red — these are the "
            "candidate overlap pairs Section D.2 enumerates."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True; r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # D.2 — Top cross-WG overlap pairs (text + table)
    overlaps = _compute_overlap_pairs(questions, similarity_data)
    if overlaps is not None and not overlaps.empty:
        doc.add_heading("D.2 Top cross-WG overlap pairs", level=2)
        doc.add_paragraph(
            "Pairs of questions in different working groups with cosine "
            "similarity ≥ 0.55, sorted by similarity. Co-leads should "
            "review each pair and pick: complementary (keep both, "
            "cross-reference), coordinate (align scope/timing), or "
            "merge (consolidate into one)."
        )
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(["WG-A · Q", "WG-B · Q", "Question texts", "Similarity"]):
            c = t.rows[0].cells[i]; c.text = h
            for r in c.paragraphs[0].runs: r.bold = True
        for _, r in overlaps.head(25).iterrows():
            row = t.add_row().cells
            row[0].text = f"WG{int(r['wg_a'])} · Q{int(r['qid_a'])}"
            row[1].text = f"WG{int(r['wg_b'])} · Q{int(r['qid_b'])}"
            short_a = (r['text_a'] or "").replace("\n", " ")
            short_b = (r['text_b'] or "").replace("\n", " ")
            if len(short_a) > 100: short_a = short_a[:99] + "…"
            if len(short_b) > 100: short_b = short_b[:99] + "…"
            row[2].text = f"A: {short_a}\nB: {short_b}"
            row[3].text = f"{r['similarity']:.2f}"

    # F7 — theme dendrogram
    if "F7" in cross_figures:
        doc.add_heading("D.3 Theme clusters across the agenda", level=2)
        doc.add_picture(io.BytesIO(cross_figures["F7"]), width=Inches(6.6))
        cap = doc.add_paragraph(
            "Figure D.3 — Hierarchical (Ward) clustering of every active "
            "question across all WGs, with Opus-generated cluster labels "
            "below."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True; r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    if theme_clusters:
        doc.add_heading("D.3a Cluster labels", level=3)
        for c in theme_clusters:
            p = doc.add_paragraph(style="List Bullet")
            label = c.get("label") or "—"
            desc = c.get("description") or ""
            run = p.add_run(label)
            run.bold = True
            p.add_run(f" — {desc}")

    # F8 — pillar coverage
    if "F8" in cross_figures:
        doc.add_heading("D.4 Pillar coverage", level=2)
        doc.add_picture(io.BytesIO(cross_figures["F8"]), width=Inches(5.5))
        cap = doc.add_paragraph(
            "Figure D.4 — Distribution of questions across the four "
            "Technology / Training / Self / Society pillars by working "
            "group. (n×) marks how many of each cell's questions were "
            "tagged cross-cutting (touching ≥3 pillars)."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True; r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # F9 — cross-cutting heatmap
    if "F9" in cross_figures:
        doc.add_heading("D.5 Cross-cutting topics", level=2)
        doc.add_picture(io.BytesIO(cross_figures["F9"]), width=Inches(6.4))
        cap = doc.add_paragraph(
            "Figure D.5 — Distribution of cross-cutting topics by working "
            "group. Tags assigned by Claude Opus from a fixed 10-tag set; "
            "spot-check the audit log if a count looks surprising."
        )
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True; r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_page_break()


def _compute_overlap_pairs(
    questions: pd.DataFrame,
    sim_data: Optional[tuple] = None,
) -> Optional[pd.DataFrame]:
    """If we have similarity data, compute the cross-WG pair table.
    Currently called only when cross_figures includes the matrix as
    `_sim_data`; the orchestrator passes it through.
    """
    if sim_data is None or questions.empty:
        return None
    qids, sim = sim_data
    qmap = questions.set_index("question_id")
    rows = []
    for i in range(len(qids)):
        for j in range(i + 1, len(qids)):
            s = float(sim[i, j])
            if s < 0.55:
                continue
            qa, qb = qids[i], qids[j]
            wga = int(qmap.loc[qa, "wg_id"])
            wgb = int(qmap.loc[qb, "wg_id"])
            if wga == wgb:
                continue
            rows.append({
                "qid_a": qa, "qid_b": qb,
                "wg_a": wga, "wg_b": wgb,
                "text_a": qmap.loc[qa, "text"],
                "text_b": qmap.loc[qb, "text"],
                "similarity": s,
            })
    df = pd.DataFrame(rows).sort_values("similarity", ascending=False).reset_index(drop=True)
    return df


def _section_e_stub(doc: "Document") -> None:
    doc.add_heading("E. Round 2 plan", level=1)
    p = doc.add_paragraph()
    p.add_run("[Pending — Phase 5] ").italic = True
    p.add_run(
        "Auto-drafted recommendations per question for Round 2, marked for "
        "co-lead review. Includes: which questions go forward unchanged, "
        "which are revised (and how), and which retire."
    )
    doc.add_page_break()


def _appendix_stub(doc: "Document") -> None:
    doc.add_heading("Appendices", level=1)
    doc.add_heading("Appendix 1 — Full per-WG question list", level=2)
    doc.add_paragraph("[Pending — long-format table of every question with all six metrics, bucket, and a co-lead disposition column for Round 2.]")
    doc.add_heading("Appendix 2 — All Round 1 free-text comments by question (anonymous)", level=2)
    doc.add_paragraph("[Pending]")
    doc.add_heading("Appendix 3 — AI synthesis audit log", level=2)
    doc.add_paragraph("[Pending]")


# --- Helpers --------------------------------------------------------------

def _pct(num, denom) -> str:
    if not denom:
        return "—"
    return f"{(num / denom) * 100:.0f}%"
